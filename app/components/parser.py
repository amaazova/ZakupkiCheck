"""Document parser: docx/pdf/txt/md → plain text."""
from __future__ import annotations

import io
from typing import Any


def _parse_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
        parts.append("")

    return "\n".join(parts).strip()


def _parse_pdf(data: bytes) -> str:
    import fitz

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as pdf:
        for page in pdf:
            text = page.get_text("text")
            if text:
                parts.append(text.strip())
    return "\n\n".join(parts).strip()


def _parse_text(data: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def parse_document(uploaded_file: Any) -> str:
    name = getattr(uploaded_file, "name", "") or ""
    data = uploaded_file.read()
    if isinstance(data, str):
        data = data.encode("utf-8")

    suffix = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if suffix == "docx":
        return _parse_docx(data)
    if suffix == "pdf":
        return _parse_pdf(data)
    if suffix in ("txt", "md"):
        return _parse_text(data)
    return _parse_text(data)
